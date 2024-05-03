from docx import Document


def replace_text(doc, old_text, new_text):

    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:

                if old_text in run.text:

                    run.text = run.text.replace(old_text, new_text)


def addInfo(list, template):
    for student in list:
        doc = Document(template)
        replace_text(doc, "Name", student.name)
        replace_text(doc, "CLASS", student.sem + student.dept)
        replace_text(doc, "100MSPRINTANDJUMP", student.event1)

        # Save the modified document
        doc.save(f"Final_Word2//{student.name}_certificate.docx")

        if student.event2 != "":

            doc = Document(template)
            replace_text(doc, "Name", student.name)
            replace_text(doc, "CLASS", student.sem + student.dept)
            replace_text(doc, "100MSPRINTANDJUMP", student.event2)
            doc.save(f"Final_Word//{student.name}_certificate2.docx")

# doc=Document("template.docx")
# for paragraph in doc.paragraphs:
#     print(paragraph.text)